# -*- coding: utf-8 -*-
# PySide6 GUI: мультивыбор файлов → локальная транскрибация (faster-whisper) с прогрессом → summary (OpenAI) → DOCX рядом с видео
# Зависимости: pip install pyside6 faster-whisper python-docx openai httpx ffmpeg-python

import os
import time
from dataclasses import dataclass
from typing import List

# (опционально) пути к CUDA/CuDNN (Windows)
try:
    os.add_dll_directory(r"C:\Program Files\NVIDIA GPU Computing Toolkit\CUDA\v12.6\bin")
    os.add_dll_directory(r"C:\Program Files\NVIDIA\CUDNN\v9.12\bin\12.9")
except Exception:
    pass

# --- OpenAI (summary) ---
import openai
import httpx
from openai.types.chat import ChatCompletionUserMessageParam

# --- STT (локально) ---
from faster_whisper import WhisperModel

# --- DOCX ---
from docx import Document

# --- Qt ---
from PySide6.QtCore import QObject, Signal, QThread, QTimer, Qt
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QLabel,
    QPushButton, QProgressBar, QFileDialog, QMessageBox
)

# ================== НАСТРОЙКИ ==================
# OpenAI
openai.api_key = os.getenv("OPEN_API_KEY")   # ключ храните в переменной окружения
PROXY_URL = os.getenv("PROXY")
_oai_client = openai.OpenAI(api_key=openai.api_key, http_client=httpx.Client(proxy=PROXY_URL))

# STT (локально)
LOCAL_MODEL_NAME = "large-v3"                     # при OOM будет fallback
COMPUTE_TRY = ["float16", "int8_float16", "int8"] # порядок попыток на GPU
CHUNK_LENGTH = 30
USE_VAD = True

# Chat-модель для summary (по индексу)
TEXT_MODELS = ["gpt-5","gpt-5-mini","gpt-4.1","gpt-4.1-mini","gpt-4o","gpt-4o-mini"]
MODEL_INDEX = 5  # gpt-4o-mini

# =================================================

def _supports_temperature(model_name: str) -> bool:
    return ("mini" in model_name) or (model_name == "gpt-4o")


def make_summary(transcript_text: str) -> str:
    if not openai.api_key:
        return "Summary отключён: OPENAI_API_KEY не задан."
    prompt = (
        "Ты выступаешь как карьерный консультант. "
        "Из транскрипта собеседования выдели список часто задаваемых вопросов "
        "и краткие корректные ответы. Формат: маркированный список, "
        "вопрос — **жирным**, ответ — обычным.\n\n"
        f"{transcript_text}"
    )
    model = TEXT_MODELS[MODEL_INDEX]
    kwargs = dict(
        model=model,
        messages=[ChatCompletionUserMessageParam(role="user", content=prompt)],
    )
    if _supports_temperature(model):
        kwargs["temperature"] = 0.2

    try:
        resp = _oai_client.chat.completions.create(**kwargs)
        msg = resp.choices[0].message
        if isinstance(msg.content, list):
            parts = []
            for part in msg.content:
                if getattr(part, "type", None) == "text":
                    parts.append(getattr(part, "text", ""))
            return "\n".join(parts).strip()
        return str(msg.content).strip()
    except Exception as e:
        return f"Summary не создано: {e}"


# -------- Whisper загрузка (singleton) --------
_whisper_model = None
def get_local_model() -> WhisperModel:
    global _whisper_model
    if _whisper_model is not None:
        return _whisper_model

    last_err = None
    # GPU
    for ct in COMPUTE_TRY:
        try:
            _whisper_model = WhisperModel(LOCAL_MODEL_NAME, device="cuda", compute_type=ct)
            return _whisper_model
        except Exception as e:
            last_err = e
    # CPU fallback
    try:
        _whisper_model = WhisperModel(LOCAL_MODEL_NAME, device="cpu", compute_type="int8")
        return _whisper_model
    except Exception as e:
        raise RuntimeError(f"Не удалось инициализировать Whisper: {e or last_err}")


# --------- Модель данных ----------
@dataclass
class JobResult:
    out_path: str
    elapsed: str


# --------- Worker в QThread ----------
class TranscribeWorker(QObject):
    progress = Signal(float, float)          # (current_sec, total_sec)
    file_started = Signal(str)               # file path
    finished = Signal(JobResult)             # результат
    failed = Signal(str)                     # сообщение об ошибке
    log5min = Signal(int)                    # каждые 5 минут: прошедшие минуты

    def __init__(self, file_path: str):
        super().__init__()
        self.file_path = file_path
        self._stop_timer = False

    def _start_5min_logger(self):
        self._stop_timer = False
        self._start_time = time.time()
        self._timer = QTimer()
        self._timer.setInterval(300000)  # 5 минут
        self._timer.timeout.connect(self._on_5min_tick)
        self._timer.start()

    def _stop_5min_logger(self):
        self._stop_timer = True
        if hasattr(self, "_timer"):
            self._timer.stop()

    def _on_5min_tick(self):
        elapsed = int((time.time() - self._start_time) // 300) * 5
        if not self._stop_timer:
            self.log5min.emit(max(5, elapsed))

    def run(self):
        try:
            self.file_started.emit(self.file_path)
            self._start_5min_logger()
            t0 = time.time()

            model = get_local_model()

            text_parts: List[str] = []
            segments, info = model.transcribe(
                self.file_path,
                vad_filter=USE_VAD,
                language=None,
                chunk_length=CHUNK_LENGTH
            )

            total = float(getattr(info, "duration", 0.0) or 0.0)
            self.progress.emit(0.0, total)

            for seg in segments:
                text_parts.append(seg.text)
                if seg.end is not None:
                    self.progress.emit(float(seg.end), total)

            transcript = "".join(text_parts).strip()

            self._stop_5min_logger()

            # Summary
            summary = make_summary(transcript)

            # DOCX
            base, _ = os.path.splitext(self.file_path)
            out_path = base + "_interview.docx"
            doc = Document()
            doc.add_heading("Транскрипция собеседования", level=1)
            doc.add_paragraph(transcript)
            doc.add_heading("Summary: Частые вопросы и ответы", level=1)
            for line in summary.splitlines():
                doc.add_paragraph(line)
            doc.save(out_path)

            elapsed = int(time.time() - t0)
            m, s = divmod(elapsed, 60); h, m = divmod(m, 60)
            self.finished.emit(JobResult(out_path=out_path, elapsed=f"{h:02d}:{m:02d}:{s:02d}"))

        except Exception as e:
            self._stop_5min_logger()
            self.failed.emit(str(e))


# --------------- GUI ---------------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Interview Transcriber (PySide6)")
        self.setFixedSize(560, 240)

        central = QWidget(self)
        layout = QVBoxLayout(central)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)

        self.lbl_title = QLabel("Транскрибация + Summary (локально + ChatGPT)")
        self.lbl_title.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        self.lbl_title.setStyleSheet("font-size:16px; font-weight:600;")
        layout.addWidget(self.lbl_title)

        self.btn_choose = QPushButton("Выбрать файлы…")
        self.btn_choose.clicked.connect(self.choose_files)
        layout.addWidget(self.btn_choose)

        self.btn_start = QPushButton("Старт")
        self.btn_start.setEnabled(False)
        self.btn_start.clicked.connect(self.start_queue)
        layout.addWidget(self.btn_start)

        self.lbl_file = QLabel("Файл: —")
        self.lbl_file.setStyleSheet("color:#444;")
        layout.addWidget(self.lbl_file)

        self.pb = QProgressBar()
        self.pb.setRange(0, 100)
        self.pb.setValue(0)
        layout.addWidget(self.pb)

        self.lbl_stats = QLabel("0%  •  00:00 / --:--  •  ETA --:--")
        self.lbl_stats.setStyleSheet("color:#666;")
        layout.addWidget(self.lbl_stats)

        self.setCentralWidget(central)

        self.queue: List[str] = []
        self.current_thread: QThread | None = None
        self.current_worker: TranscribeWorker | None = None
        self.total_secs = 0.0
        self.start_time = 0.0

    def choose_files(self):
        paths, _ = QFileDialog.getOpenFileNames(
            self, "Выберите видео/аудио", "",
            "Видео/аудио (*.mp4 *.mkv *.mp3 *.wav *.m4a *.mov *.avi)"
        )
        self.queue = list(paths or [])
        self.btn_start.setEnabled(bool(self.queue))
        self.lbl_file.setText(f"Выбрано файлов: {len(self.queue)}" if self.queue else "Файл: —")
        self.pb.setValue(0)
        self.lbl_stats.setText("0%  •  00:00 / --:--  •  ETA --:--")

    def start_queue(self):
        if not self.queue:
            return
        self.btn_choose.setEnabled(False)
        self.btn_start.setEnabled(False)
        self.process_next()

    def process_next(self):
        if not self.queue:
            self.btn_choose.setEnabled(True)
            self.btn_start.setEnabled(True)
            QMessageBox.information(self, "Готово", "Все файлы обработаны.")
            return

        file_path = self.queue.pop(0)
        self.run_job(file_path)

    def run_job(self, file_path: str):
        # подготовка UI
        self.lbl_file.setText(f"Файл: {os.path.basename(file_path)}")
        self.pb.setRange(0, 100)
        self.pb.setValue(0)
        self.lbl_stats.setText("0%  •  00:00 / --:--  •  ETA --:--")
        self.start_time = time.time()
        self.total_secs = 0.0

        # запуск worker в QThread
        self.current_thread = QThread()
        self.current_worker = TranscribeWorker(file_path)
        self.current_worker.moveToThread(self.current_thread)

        # сигналы
        self.current_thread.started.connect(self.current_worker.run)
        self.current_worker.file_started.connect(lambda p: print(f"[INFO] Старт транскрибации: {p}"))
        self.current_worker.log5min.connect(lambda mins: print(f"[INFO] Прошло {mins} минут..."))
        self.current_worker.progress.connect(self.on_progress)
        self.current_worker.finished.connect(self.on_finished)
        self.current_worker.failed.connect(self.on_failed)

        # очистка
        self.current_worker.finished.connect(self.cleanup_thread)
        self.current_worker.failed.connect(self.cleanup_thread)

        self.current_thread.start()

    def on_progress(self, current: float, total: float):
        self.total_secs = total
        pct = int(100 * min(current, total or 1) / (total or 1))
        self.pb.setValue(pct)

        elapsed = time.time() - self.start_time
        eta = (total - current) * (elapsed / current) if total and current > 0 else 0

        def fmt(t: float) -> str:
            t = int(max(0, t))
            m, s = divmod(t, 60); h, m = divmod(m, 60)
            return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"

        self.lbl_stats.setText(f"{pct}%  •  {fmt(current)} / {fmt(total)}  •  ETA {fmt(eta)}")

    def on_finished(self, result: JobResult):
        print(f"[INFO] Готово: {result.out_path}. Общее время: {result.elapsed}")
        QMessageBox.information(self, "Готово", f"Файл сохранён:\n{result.out_path}\n\nОбщее время: {result.elapsed}")
        self.process_next()

    def on_failed(self, message: str):
        QMessageBox.critical(self, "Ошибка", message)
        self.process_next()

    def cleanup_thread(self):
        if self.current_thread:
            self.current_thread.quit()
            self.current_thread.wait()
        self.current_thread = None
        self.current_worker = None


if __name__ == "__main__":
    import sys
    app = QApplication(sys.argv)
    w = MainWindow()
    w.show()
    sys.exit(app.exec())
